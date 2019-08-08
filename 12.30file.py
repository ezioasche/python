# -*- coding: utf-8 -*-
"""
需求：把一堆.JAVA文件中的文本拿出来并保存到word中（做软著）
"""
from docx import Document  # 百度得到这个有趣强大的库
import os
#######以下三个变量分别是根目录路径、需要读取的文件种类、生成word的文件路径，自己填########
path = 'E:\\JAVA\\Gitproject\\sign_page\\signSystem-master\\signSystem-master\\src\\main\\java'
filetype = 'java'
savepath = 'E:\\program\\Gitcopy\\python\\signSystem.docx'
################

document = Document()

linesnum = 0
filepath = []
for root, dirs, files in os.walk(path):

    for file in files:

        if filetype in file:
            document.add_heading(file)  # 以文件名为小标题
            file = os.path.join(root, file)  # 标题写好再得到路径
            f = open(file, encoding="utf-8")
            linesnum += len(f.readlines())
            f.seek(0)
            nr = f.read()
            paragraph = document.add_paragraph(nr)
            # data.append(nr)
            f.close()

            filepath.append(file)

document.save(savepath)

print(filepath)  # 文件中所有.java结尾的文件
print("---java文件总行数为："+str(linesnum)+" ---")
